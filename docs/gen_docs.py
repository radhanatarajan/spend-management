"""
Generate spend_analytics_guide.docx and contract_database_guide.docx
matching the style of spend_management_setup.docx.

Run:  uv run --with python-docx python3 docs/gen_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

# ── Colour palette (matched from spend_management_setup.docx) ─────────────────
DARK_BLUE  = RGBColor(0x1E, 0x5F, 0xA5)
MED_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
DARK_GREY  = RGBColor(0x40, 0x40, 0x40)
BLACK      = RGBColor(0x00, 0x00, 0x00)
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

TODAY = date.today().strftime("%B %d, %Y")


# ── Builder helpers ───────────────────────────────────────────────────────────

def _run(para, text, bold=False, italic=False, size_pt=11, color=BLACK, mono=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    if mono:
        run.font.name = "Courier New"
    return run


def doc_title(doc, title, subtitle):
    p = doc.add_paragraph()
    _run(p, title, bold=True, size_pt=26, color=DARK_BLUE)
    p2 = doc.add_paragraph()
    _run(p2, subtitle, bold=True, size_pt=18, color=MED_BLUE)
    p3 = doc.add_paragraph()
    _run(p3, f"Date: {TODAY}", size_pt=11, color=DARK_GREY)
    doc.add_paragraph()


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(16)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = MED_BLUE
        run.font.size = Pt(13)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = DARK_GREY
        run.font.size = Pt(12)
    return p


def body(doc, text):
    p = doc.add_paragraph()
    _run(p, text, size_pt=11, color=BLACK)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    _run(p, text, size_pt=11, color=BLACK)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    for line in text.strip().split("\n"):
        run = p.add_run(line + "\n")
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK_GREY
    return p


def _set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # Header row
    hdr_row = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        _set_cell_bg(cell, "1E5FA5")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(10)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        bg = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(cell_text)
            _set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# ══════════════════════════════════════════════════════════════════════════════
# SPEND ANALYTICS GUIDE
# ══════════════════════════════════════════════════════════════════════════════

def build_spend_doc():
    doc = Document()

    doc_title(doc,
        "Spend Analytics Module",
        "Feature Guide — Spend Management Platform")

    # ── 1. Overview ──────────────────────────────────────────────────────────
    h1(doc, "1. Overview")
    body(doc,
        "The Spend Analytics module is the primary data-exploration surface of the platform. "
        "It presents all Oracle-aligned General Ledger (GL) spend transactions in a fully "
        "interactive table with cross-dependent slicers, multi-column sorting, pagination, "
        "KPI summary cards, and a dedicated Spend Report page with charts and CSV export. "
        "All data is stored in MySQL and served through a paginated FastAPI endpoint.")

    h2(doc, "1.1 Navigation")
    bullet(doc, "Sidebar link: Spend Analytics → /spend")
    bullet(doc, "Reports section → Spend Report → /reports/spend")

    h2(doc, "1.2 Seed data")
    body(doc,
        "On first startup the backend seeds ~150 synthetic spend transactions "
        "across 8 vendors, multiple departments, and the last 12 months. "
        "Data is seeded only once — if rows already exist the seed is skipped.")

    # ── 2. Spend Analytics page ───────────────────────────────────────────────
    h1(doc, "2. Spend Analytics Page (/spend)")

    h2(doc, "2.1 KPI Summary Strip")
    body(doc,
        "Four metric cards sit above the table and update in real time as filters change:")
    table(doc,
        ["Card", "Metric"],
        [
            ["Total Spend", "Sum of amount_usd across all filtered rows"],
            ["Transactions", "Row count after filters are applied"],
            ["Avg per Month", "Total spend ÷ distinct months in the filtered result"],
            ["Top Vendor", "Vendor name with the highest aggregate spend"],
        ],
        col_widths=[2.0, 4.5])

    h2(doc, "2.2 Slicer Filters")
    body(doc,
        "Seven dropdown slicers appear above the table. Each slicer shows only the "
        "values present in the data after the other slicers are applied "
        "(cross-dependent filtering). Selecting a value triggers a backend query that "
        "re-fetches both the transaction list and the available filter options.")
    table(doc,
        ["Slicer", "Field", "Notes"],
        [
            ["Month",        "month_key",             "YYYYMM integer; displayed as 'Jan '25'"],
            ["Expense Type", "expense_type",          "Capex, Opex, Travel, etc."],
            ["Company Code", "company_code",          "Oracle company identifier"],
            ["Department",   "oracle_department",     "Oracle dept code + name"],
            ["Acct Group",   "oracle_account_group",  "R&D, S&M, G&A, Infra, …"],
            ["Vendor",       "vendor_name",           "Supplier name"],
            ["JE Source",    "je_source",             "Coupa, Concur, Workday, Oracle, Manual"],
        ],
        col_widths=[1.4, 2.0, 3.1])
    bullet(doc, "Select one or multiple values per slicer.")
    bullet(doc, "Click Apply Filters to update the table.")
    bullet(doc, "Click Clear inside a slicer to reset only that filter.")
    bullet(doc, "Click Reset All to clear every filter simultaneously.")

    h2(doc, "2.3 Transaction Table")
    body(doc,
        "The table shows up to 50 rows per page by default. "
        "Each column header is clickable for ascending/descending sort. "
        "The header uses a two-row grouped layout to cluster related Oracle columns.")
    table(doc,
        ["Column", "Field", "Notes"],
        [
            ["Month",            "month_label",                    "Derived from month_key"],
            ["Expense Type",     "expense_type",                   ""],
            ["Co. Code",         "company_code",                   ""],
            ["Oracle Org",       "oracle_organization",            "Business unit"],
            ["Acct No.",         "oracle_account_number",          "GL account number"],
            ["Dept Code",        "oracle_department",              ""],
            ["Dept Name",        "oracle_department_name",         ""],
            ["Hierarchy",        "oracle_cost_center_hierarchy",   "Cost center rollup"],
            ["Acct Group",       "oracle_account_group",           ""],
            ["Acct Sub-Group",   "oracle_account_sub_group",       ""],
            ["Cost Element",     "oracle_cost_element",            "Salaries, Licenses, …"],
            ["Line Desc",        "line_desc",                      "Free-text description"],
            ["Vendor",           "vendor_name",                    ""],
            ["PO",               "po_recon",                       "Hover for description"],
            ["PO Number",        "purchase_order_number",          ""],
            ["PO Line",          "purchase_order_line_number",     ""],
            ["Invoice No.",      "invoice_number",                 ""],
            ["Invoice Line",     "invoice_line_number",            ""],
            ["JE Source",        "je_source",                      ""],
            ["$",                "amount_usd",                     "USD, right-aligned"],
        ],
        col_widths=[1.4, 1.8, 3.3])

    h2(doc, "2.4 Pagination")
    bullet(doc, "Default page size: 50 rows.")
    bullet(doc, "Page controls appear at the bottom of the table.")
    bullet(doc, "Total row count is shown in the summary strip.")

    # ── 3. Spend Report page ──────────────────────────────────────────────────
    h1(doc, "3. Spend Report Page (/reports/spend)")
    body(doc,
        "The Spend Report page provides an executive-level view of spend with KPI cards, "
        "charts, period comparison, and CSV export. It uses a fiscal period picker at the top right.")

    h2(doc, "3.1 Financial Period Picker")
    body(doc,
        "Pill buttons at the top-right let the user select the reporting window. "
        "The selected period filters all KPIs, charts, and the insights panel.")

    h2(doc, "3.2 KPI Cards")
    table(doc,
        ["Card", "Description"],
        [
            ["Total Spend",    "Sum of all transactions in the selected period"],
            ["Transactions",   "Count of GL lines in the period"],
            ["Top Vendor",     "Vendor with highest aggregate spend"],
            ["Avg / Month",    "Total spend divided by months in the period"],
        ],
        col_widths=[2.0, 4.5])

    h2(doc, "3.3 Charts")
    bullet(doc, "Monthly Spend Trend — bar chart of spend per month in the period.")
    bullet(doc, "Spend by Account Group — donut chart (R&D, S&M, G&A, Infra, …).")
    bullet(doc, "Top Vendors by Spend — horizontal bar chart, top 8 vendors.")
    bullet(doc, "Spend by Department — breakdown across Oracle departments.")

    h2(doc, "3.4 Insights Panel")
    body(doc,
        "Auto-generated text insights highlight the largest spend category, "
        "the fastest-growing vendor, and period-over-period change.")

    h2(doc, "3.5 CSV Export")
    body(doc,
        "The Download CSV button exports all transactions matching the current "
        "filter to a comma-separated file. The file includes all 20 GL columns.")

    # ── 4. API Reference ──────────────────────────────────────────────────────
    h1(doc, "4. API Reference")
    body(doc, "Base URL: http://localhost:8000   •   Swagger UI: http://localhost:8000/docs")

    table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/api/spend/transactions",   "Paginated, filtered, sorted transaction list"],
            ["GET", "/api/spend/filter-options", "Distinct slicer values (cross-filtered by active filters)"],
        ],
        col_widths=[0.8, 2.8, 3.0])

    h2(doc, "4.1 GET /api/spend/transactions — Query Parameters")
    table(doc,
        ["Parameter", "Type", "Default", "Description"],
        [
            ["page",                    "int",    "1",    "Page number (1-indexed)"],
            ["page_size",               "int",    "50",   "Rows per page"],
            ["sort_by",                 "string", "month_key", "Column to sort by"],
            ["sort_dir",                "string", "desc", "asc or desc"],
            ["month_key",               "int[]",  "—",    "Filter by one or more YYYYMM keys"],
            ["expense_type",            "str[]",  "—",    "Filter by expense type(s)"],
            ["company_code",            "str[]",  "—",    "Filter by company code(s)"],
            ["oracle_department",       "str[]",  "—",    "Filter by department code(s)"],
            ["oracle_account_group",    "str[]",  "—",    "Filter by account group(s)"],
            ["vendor_name",             "str[]",  "—",    "Filter by vendor name(s)"],
            ["je_source",               "str[]",  "—",    "Filter by journal entry source(s)"],
        ],
        col_widths=[1.8, 0.8, 0.9, 3.0])

    h2(doc, "4.2 GET /api/spend/filter-options — Query Parameters")
    body(doc,
        "Accepts the same filter parameters as /transactions. "
        "Returns distinct values for each slicer dimension, constrained to rows that "
        "match the currently applied filters (cross-dependent filtering).")

    # ── 5. Data model ─────────────────────────────────────────────────────────
    h1(doc, "5. Data Model")

    h2(doc, "5.1 SpendTransaction table")
    table(doc,
        ["Column", "Type", "Notes"],
        [
            ["id",                           "INT PK",       "Auto-increment"],
            ["month_key",                    "INT",          "YYYYMM format, e.g. 202601"],
            ["month_label",                  "VARCHAR",      "e.g. 'Jan 2026'"],
            ["expense_type",                 "VARCHAR",      "Capex / Opex / Travel / …"],
            ["company_code",                 "VARCHAR",      "Oracle company code"],
            ["oracle_organization",          "VARCHAR",      "Business unit"],
            ["oracle_account_number",        "VARCHAR",      "GL account number"],
            ["oracle_department",            "VARCHAR",      "Department code"],
            ["oracle_department_name",       "VARCHAR",      "Department name"],
            ["oracle_cost_center_hierarchy", "VARCHAR",      "Rollup path"],
            ["oracle_account_group",         "VARCHAR",      "R&D / S&M / G&A / …"],
            ["oracle_account_sub_group",     "VARCHAR",      "Sub-classification"],
            ["oracle_cost_element",          "VARCHAR",      "Salaries / Licenses / …"],
            ["line_desc",                    "VARCHAR",      "Free-text description"],
            ["vendor_name",                  "VARCHAR",      "Supplier name"],
            ["po_recon",                     "VARCHAR",      "PO reconciliation code"],
            ["purchase_order_number",        "VARCHAR",      ""],
            ["purchase_order_line_number",   "VARCHAR",      ""],
            ["invoice_number",               "VARCHAR",      ""],
            ["invoice_line_number",          "VARCHAR",      ""],
            ["je_source",                    "VARCHAR",      "Journal entry source"],
            ["amount_usd",                   "DECIMAL(14,2)","Transaction amount in USD"],
        ],
        col_widths=[2.2, 1.2, 3.1])

    doc.save("docs/spend_analytics_guide.docx")
    print("✓ docs/spend_analytics_guide.docx written")


# ══════════════════════════════════════════════════════════════════════════════
# CONTRACT DATABASE GUIDE
# ══════════════════════════════════════════════════════════════════════════════

def build_contracts_doc():
    doc = Document()

    doc_title(doc,
        "Contract Database Module",
        "Feature Guide — Spend Management Platform")

    # ── 1. Overview ──────────────────────────────────────────────────────────
    h1(doc, "1. Overview")
    body(doc,
        "The Contract Database module tracks multi-year software license agreements. "
        "Each contract represents one Purchase Order (PO). A contract has one or more "
        "PO lines, where each line covers a service period (typically one year) with its "
        "own billing amount and interval. The module supports full CRUD on contracts and "
        "lines, billing interval conversion, multi-year detection, 100% renewal assumption "
        "for future periods, and a monthly breakout Contract Report.")

    h2(doc, "1.1 Navigation")
    bullet(doc, "Sidebar: Contract Database → /contracts")
    bullet(doc, "Reports section → Contract Report → /reports/contracts")

    h2(doc, "1.2 Core concepts")
    table(doc,
        ["Concept", "Description"],
        [
            ["Contract",        "One PO header: vendor, department, account, PO number, status"],
            ["Contract Line",   "One service period inside a PO: dates, billing interval, amount"],
            ["Multi-year",      "A PO with ≥2 consecutive lines (each line's end + 1 month = next line's start)"],
            ["Monthly Amount",  "Canonical per-month value stored on every line, derived from entered_amount + interval"],
            ["Renewal Assumption", "Months beyond the last signed PO line extend that line's rate at 100%, shown in amber"],
        ],
        col_widths=[1.8, 4.7])

    h2(doc, "1.3 Seed data")
    body(doc, "Four contracts are seeded automatically on first startup (skipped if data already exists):")
    table(doc,
        ["Vendor", "PO Number", "Interval", "Lines", "Status"],
        [
            ["Salesforce", "PO-83353", "Monthly",       "3 (May-26→Apr-29)", "Active"],
            ["GitHub",     "PO-91200", "Quarterly",     "2 (Jan-26→Dec-27)", "Active"],
            ["Workday",    "PO-77040", "Yearly",        "1 (Apr-26→Mar-27)", "Active"],
            ["Tableau",    "PO-60301", "Custom (total)","3 (Jan-23→Dec-25)", "Expired"],
        ],
        col_widths=[1.4, 1.2, 1.3, 2.2, 1.0])

    # ── 2. Contract Database page ─────────────────────────────────────────────
    h1(doc, "2. Contract Database Page (/contracts)")

    h2(doc, "2.1 Contract Grouping")
    body(doc,
        "Contracts are visually grouped by the combination of Vendor + Department Code + "
        "Account Number + PO Number. If multiple Contract DB records share the same key "
        "(e.g. from a data import), their lines are merged under one row. "
        "The group key is computed on the client — no backend change is required.")
    bullet(doc, "Multi-year groups display an indigo 'Multi-year' badge next to the vendor name.")
    bullet(doc, "The Terms column shows the total number of PO lines (e.g. '3 years').")
    bullet(doc, "Contract Value shows the sum of all line period totals.")
    bullet(doc, "Status reflects the contract that owns the last (most recent) line.")

    h2(doc, "2.2 Status Filter")
    body(doc,
        "Pill buttons at the top filter the table by contract status: "
        "All · Active · Pending · Expired · Cancelled. "
        "The summary strip at the right shows the count of visible POs and total value.")

    h2(doc, "2.3 Expanding a Contract Group")
    body(doc,
        "Click any row to expand it. The expanded panel shows:")
    bullet(doc, "Optional contract description.")
    bullet(doc, "A lines sub-table with columns: PO Line, Service Period, Months, Interval, Entered Amount, Monthly, Period Total, actions.")
    bullet(doc, "An Add line form at the bottom.")

    h2(doc, "2.4 Editing Contract Lines (Inline)")
    body(doc,
        "Hover over any line row to reveal Edit and ✕ (delete) buttons. "
        "Clicking Edit transforms the entire row into an inline form:")
    table(doc,
        ["Field", "Input Type", "Notes"],
        [
            ["PO Line #",     "Number",   "Integer ≥ 1"],
            ["Service Start", "Date",     "ISO date picker"],
            ["Service End",   "Date",     "ISO date picker"],
            ["Interval",      "Dropdown", "Monthly / Quarterly / Annual / Custom (total)"],
            ["Entered Amount","Text",     "The amount as entered under the chosen interval"],
            ["Notes",         "Text",     "Optional free-text note"],
        ],
        col_widths=[1.5, 1.2, 3.8])
    bullet(doc, "Months is computed live from the date fields.")
    bullet(doc, "Monthly amount preview updates as you type the amount or change the interval.")
    bullet(doc, "Press Enter or click Save to commit. Press Escape or Cancel to discard.")
    bullet(doc, "On save the backend recomputes monthly_amount and stores both values.")

    h2(doc, "2.5 Editing Contract Header")
    body(doc,
        "Click the Edit button (visible in the rightmost column of each group row) "
        "to open the Edit Contract modal. It pre-fills all header fields:")
    bullet(doc, "Vendor Name, Department Code, Department Name")
    bullet(doc, "Account Number, Account Sub-Group, PO Number")
    bullet(doc, "Status (Active / Pending / Expired / Cancelled)")
    bullet(doc, "Description")
    body(doc,
        "If the group spans multiple Contract DB records (merged group), "
        "the same header update is applied to all records simultaneously.")

    h2(doc, "2.6 Adding a New Contract")
    body(doc,
        "Click + New Contract (top-right). The modal accepts all header fields plus "
        "one or more PO lines. Click + Add line to append additional lines before saving. "
        "Monthly amount preview appears for non-monthly intervals.")

    h2(doc, "2.7 Deleting")
    bullet(doc, "Delete button on a group row: deletes all Contract DB records in the group after confirmation.")
    bullet(doc, "✕ on a line: removes only that line from the contract.")

    # ── 3. Billing intervals ──────────────────────────────────────────────────
    h1(doc, "3. Billing Intervals")
    body(doc,
        "Every PO line has a billing_interval that determines how the entered_amount "
        "maps to a monthly_amount. Both values are stored; monthly_amount is always "
        "the canonical per-month figure used in reports.")
    table(doc,
        ["Interval", "Entered Amount Meaning", "Monthly Amount Formula"],
        [
            ["Monthly",          "Amount charged per month",          "= entered_amount"],
            ["Quarterly",        "Amount charged per quarter",         "= entered_amount ÷ 3"],
            ["Annual",           "Amount charged per year",            "= entered_amount ÷ 12"],
            ["Custom (total)",   "Total amount for the full period",   "= entered_amount ÷ months_in_period"],
        ],
        col_widths=[1.4, 2.4, 2.7])
    body(doc,
        "months_in_period is computed as "
        "(end.year − start.year) × 12 + (end.month − start.month) + 1.")

    # ── 4. Multi-year detection ───────────────────────────────────────────────
    h1(doc, "4. Multi-Year Detection")
    body(doc,
        "A contract is classified as multi-year when it has at least two PO lines "
        "whose service periods are consecutive — i.e. the month immediately after "
        "line N ends is the month line N+1 starts. This check is performed both "
        "server-side (for the Contract Report) and client-side (for the badge on the "
        "Contract Database page).")

    h2(doc, "4.1 Consecutive month rule")
    body(doc, "For adjacent lines A and B (sorted by period_start):")
    code_block(doc,
        "next_month = 1              if A.period_end.month == 12\n"
        "           else A.period_end.month + 1\n"
        "next_year  = A.period_end.year + 1  if A.period_end.month == 12\n"
        "           else A.period_end.year\n"
        "\n"
        "is_consecutive = (B.period_start.year  == next_year and\n"
        "                  B.period_start.month == next_month)")

    h2(doc, "4.2 Examples")
    table(doc,
        ["Line A End", "Line B Start", "Consecutive?"],
        [
            ["2026-04-30", "2026-05-01", "Yes"],
            ["2026-12-31", "2027-01-01", "Yes"],
            ["2026-04-30", "2026-06-01", "No (gap of 1 month)"],
            ["2026-04-30", "2026-04-01", "No (overlap)"],
        ],
        col_widths=[1.6, 1.6, 2.0])

    # ── 5. Contract Report page ───────────────────────────────────────────────
    h1(doc, "5. Contract Report Page (/reports/contracts)")
    body(doc,
        "The Contract Report shows only multi-year contracts. "
        "For each contract it displays a monthly dollar breakout across the 12 months "
        "of the selected fiscal year, plus a FY Total column.")

    h2(doc, "5.1 Fiscal Year")
    body(doc,
        "The fiscal year is a calendar year (Jan – Dec). FY 2027 covers January 2027 – "
        "December 2027. The FY picker (top-right) shows only fiscal years for which "
        "at least one multi-year contract has coverage or renewal assumption.")

    h2(doc, "5.2 Renewal Assumption")
    body(doc,
        "For months that fall after the last signed PO line the system projects forward "
        "at 100% of that line's monthly_amount. These assumed values are visually "
        "distinguished:")
    bullet(doc, "Shown in amber italic text.")
    bullet(doc, "Prefixed with ~ (e.g. ~$68.9K).")
    bullet(doc, "Tooltip: '100% renewal assumption'.")
    bullet(doc, "The Assumed (renewal) KPI card shows the total assumed spend for the FY.")

    h2(doc, "5.3 Slicers")
    body(doc,
        "Three multi-select pill slicers filter the report client-side: "
        "Vendor, Department, and Status. "
        "The monthly totals row and all KPI cards update to reflect the filtered rows.")

    h2(doc, "5.4 KPI Cards")
    table(doc,
        ["Card", "Description"],
        [
            ["Contracts",          "Count of multi-year contracts visible after slicers"],
            ["Vendors",            "Distinct vendor count"],
            ["FY Total Value",     "Sum of all monthly amounts (actual + assumed) in the FY"],
            ["Assumed (renewal)",  "Portion of FY Total that is projected (not yet signed)"],
        ],
        col_widths=[2.0, 4.5])

    h2(doc, "5.5 Monthly Table")
    body(doc,
        "Columns: Vendor (sticky), Department, Account, one column per month (12 total), FY Total. "
        "The footer row shows per-month totals and the grand total. "
        "Actual amounts are shown in dark monospace font; assumed amounts in amber italic.")

    # ── 6. API Reference ──────────────────────────────────────────────────────
    h1(doc, "6. API Reference")
    body(doc, "Base URL: http://localhost:8000   •   Swagger UI: http://localhost:8000/docs")
    body(doc, "All contract endpoints require authentication (Bearer JWT).")

    h2(doc, "6.1 Contract CRUD")
    table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET",    "/api/contracts",                         "List all contracts (with lines)"],
            ["POST",   "/api/contracts",                         "Create a new contract with optional lines"],
            ["GET",    "/api/contracts/{id}",                    "Get a single contract"],
            ["PUT",    "/api/contracts/{id}",                    "Update contract header fields"],
            ["DELETE", "/api/contracts/{id}",                    "Delete contract and all its lines"],
        ],
        col_widths=[0.8, 2.8, 3.0])

    h2(doc, "6.2 Contract Line CRUD")
    table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["POST",   "/api/contracts/{id}/lines",              "Add a line to a contract"],
            ["PUT",    "/api/contracts/{id}/lines/{line_id}",    "Update a line (recomputes monthly_amount)"],
            ["DELETE", "/api/contracts/{id}/lines/{line_id}",    "Remove a line"],
        ],
        col_widths=[0.8, 2.8, 3.0])

    h2(doc, "6.3 Contract Report")
    table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/api/contracts/report?fiscal_year={year}", "Returns multi-year contract monthly breakout for a FY"],
        ],
        col_widths=[0.8, 3.4, 2.4])

    # ── 7. Data Model ─────────────────────────────────────────────────────────
    h1(doc, "7. Data Model")

    h2(doc, "7.1 contracts table")
    table(doc,
        ["Column", "Type", "Notes"],
        [
            ["id",                      "INT PK",      "Auto-increment"],
            ["vendor_name",             "VARCHAR",     ""],
            ["description",             "TEXT",        "Optional"],
            ["oracle_department",       "VARCHAR",     "Dept code, e.g. 1200"],
            ["oracle_department_name",  "VARCHAR",     "Dept name, e.g. Sales"],
            ["oracle_account_number",   "VARCHAR",     "GL account, e.g. ACC-6385"],
            ["oracle_account_sub_group","VARCHAR",     "e.g. Software Licenses"],
            ["purchase_order_number",   "VARCHAR",     "e.g. PO-83353"],
            ["status",                  "ENUM",        "active / pending / expired / cancelled"],
            ["created_at",              "DATETIME",    "Auto-set on insert"],
            ["updated_at",              "DATETIME",    "Auto-updated on save"],
        ],
        col_widths=[2.2, 1.1, 3.2])

    h2(doc, "7.2 contract_lines table")
    table(doc,
        ["Column", "Type", "Notes"],
        [
            ["id",               "INT PK",       "Auto-increment"],
            ["contract_id",      "INT FK",       "→ contracts.id (cascade delete)"],
            ["po_line_number",   "INT",          "PO line number, e.g. 4"],
            ["period_start",     "DATE",         "First day of the service period"],
            ["period_end",       "DATE",         "Last day of the service period"],
            ["billing_interval", "ENUM",         "monthly / quarterly / yearly / custom"],
            ["entered_amount",   "DECIMAL(14,2)","Amount as entered under the chosen interval"],
            ["monthly_amount",   "DECIMAL(14,2)","Computed per-month value; recomputed on every save"],
            ["notes",            "TEXT",         "Optional free-text note"],
            ["created_at",       "DATETIME",     ""],
            ["updated_at",       "DATETIME",     ""],
        ],
        col_widths=[1.8, 1.2, 3.5])

    h2(doc, "7.3 Computed fields (not stored)")
    table(doc,
        ["Field", "Formula"],
        [
            ["months_in_period", "(end.year − start.year) × 12 + (end.month − start.month) + 1"],
            ["total_amount",     "monthly_amount × months_in_period"],
            ["contract_total",   "Sum of total_amount across all lines"],
        ],
        col_widths=[1.8, 4.7])

    doc.save("docs/contract_database_guide.docx")
    print("✓ docs/contract_database_guide.docx written")


if __name__ == "__main__":
    build_spend_doc()
    build_contracts_doc()
